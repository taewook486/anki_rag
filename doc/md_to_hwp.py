"""MD → DOCX → HWP 변환 스크립트

python-docx로 DOCX를 생성한 후 HWP COM으로 HWP로 변환합니다.
HWP COM의 HTML 열기는 EUC-KR 인코딩 문제가 있으므로
DOCX 경유 방식을 사용합니다.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
import win32com.client


# ─────────────────────────────────────────
# MD → DOCX
# ─────────────────────────────────────────

def _set_heading_style(para, level: int) -> None:
    """헤딩 스타일 설정 (h1~h4)"""
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
    try:
        para.style = style_map.get(level, "Heading 4")
    except KeyError:
        pass  # 스타일 없으면 기본값 유지


def _add_run_with_inline(para, text: str) -> None:
    """**bold**, `code` 인라인 마크업을 처리하여 run 추가"""
    # **bold** 와 `code` 인라인 파싱
    pattern = re.compile(r"\*\*(.+?)\*\*|`([^`]+)`")
    cursor = 0
    for m in pattern.finditer(text):
        # 일반 텍스트
        if m.start() > cursor:
            para.add_run(text[cursor : m.start()])
        if m.group(1) is not None:
            run = para.add_run(m.group(1))
            run.bold = True
        else:
            run = para.add_run(m.group(2))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        cursor = m.end()
    if cursor < len(text):
        para.add_run(text[cursor:])


def _add_horizontal_rule(doc: Document) -> None:
    """수평선 단락 추가"""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _parse_table(doc: Document, lines: list[str], start: int) -> int:
    """마크다운 테이블 파싱 후 DOCX 테이블 생성. 종료 인덱스 반환."""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        # 구분자 줄 스킵
        if re.match(r"^\|[-: |]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1

    if not rows:
        return start + 1

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= col_count:
                break
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_run_with_inline(p, cell_text)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True

    return i


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    """마크다운 파일을 DOCX로 변환"""
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    # 기본 폰트 설정
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    in_code_block = False
    code_lines: list[str] = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── 코드 블록 ──
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # 코드 블록 종료 → 일괄 추가
                in_code_block = False
                if code_lines:
                    para = doc.add_paragraph()
                    para.style = "No Spacing" if "No Spacing" in [s.name for s in doc.styles] else "Normal"
                    run = para.add_run("\n".join(code_lines))
                    run.font.name = "Courier New"
                    run.font.size = Pt(8)
                    # 회색 배경은 DOCX 수준에서 생략 (HWP에서 랜더링 복잡도)
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # ── 빈 줄 ──
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # ── 수평선 ──
        if re.match(r"^-{3,}$", stripped) or re.match(r"^={3,}$", stripped):
            _add_horizontal_rule(doc)
            i += 1
            continue

        # ── 헤딩 ──
        heading_m = re.match(r"^(#{1,6})\s+(.+)", stripped)
        if heading_m:
            level = len(heading_m.group(1))
            title_text = heading_m.group(2)
            para = doc.add_paragraph()
            _set_heading_style(para, min(level, 4))
            _add_run_with_inline(para, title_text)
            i += 1
            continue

        # ── 테이블 ──
        if stripped.startswith("|"):
            i = _parse_table(doc, lines, i)
            continue

        # ── 목록 ──
        list_m = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.+)", line)
        if list_m:
            indent = len(list_m.group(1))
            content = list_m.group(3)
            style_name = "List Bullet" if re.match(r"[-*+]", list_m.group(2)) else "List Number"
            try:
                para = doc.add_paragraph(style=style_name)
            except KeyError:
                para = doc.add_paragraph()
            _add_run_with_inline(para, content)
            i += 1
            continue

        # ── 일반 단락 ──
        para = doc.add_paragraph()
        _add_run_with_inline(para, stripped)
        i += 1

    doc.save(str(docx_path))
    print(f"DOCX 저장 완료: {docx_path}")


# ─────────────────────────────────────────
# DOCX → HWP (COM 자동화)
# ─────────────────────────────────────────

def docx_to_hwp(docx_path: Path, hwp_path: Path) -> None:
    """HWP COM으로 DOCX → HWP 변환"""
    hwp = win32com.client.Dispatch("HWPFrame.HwpObject")
    hwp.XHwpWindows.Item(0).Visible = True  # 창 표시 (디버깅 용이)

    try:
        hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
    except Exception:
        pass  # 모듈 등록 실패 시 무시

    # DOCX 파일 열기 (절대 경로 필요)
    abs_docx = str(docx_path.resolve())
    abs_hwp = str(hwp_path.resolve())

    print(f"HWP에서 DOCX 열기: {abs_docx}")
    # 포맷 "" = 확장자 자동 감지 (DOCX 포맷 문자열이 HWP 버전마다 다름)
    result = hwp.Open(abs_docx, "", "forceopen:true")
    print(f"Open 반환값: {result}")

    print(f"HWP 파일로 저장: {abs_hwp}")
    hwp.SaveAs(abs_hwp, "HWP", "")

    hwp.Quit()
    print("변환 완료!")


# ─────────────────────────────────────────
# 메인
# ─────────────────────────────────────────

if __name__ == "__main__":
    base = Path(__file__).parent
    md_path = base / "설계서.md"
    docx_path = base / "설계서_temp.docx"
    hwp_path = base / "설계서.hwp"

    print("=== MD → DOCX ===")
    md_to_docx(md_path, docx_path)

    print("\n=== DOCX → HWP ===")
    docx_to_hwp(docx_path, hwp_path)

    print(f"\n완료: {hwp_path}")
    print("※ 설계서_temp.docx 는 검증용으로 남겨둡니다.")
